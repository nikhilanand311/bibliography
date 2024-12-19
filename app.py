import streamlit as st
import pytesseract
from PIL import Image
import fitz  # PyMuPDF for PDF text extraction
import io
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import pandas as pd
import re
import spacy
from spacy.cli import download
from transformers import pipeline, GPT2LMHeadModel, GPT2Tokenizer

# Function to ensure the SpaCy model is downloaded
def ensure_spacy_model():
    try:
        # Try loading the SpaCy model
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        # If not found, download it
        download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")
    return nlp

# Initialize spaCy model
nlp = ensure_spacy_model()

# Initialize GPT-2 text generation pipeline
generator = pipeline("text-generation", model="gpt2")

# Define a simple summarization function (can be customized)
def simple_summarize(text):
    sentences = text.split('. ')
    summary = '. '.join(sentences[:2])  # Just take the first two sentences as a summary
    return summary

# Function to extract bibliography information using regex
def extract_bibliography_info(text):
    bibliography_info = {
        'title': '',
        'author': '',
        'year': ''
    }
    
    text = text.strip().replace('\n', ' ').replace('  ', ' ')
    
    # Attempt to match title, author, and year using improved regex patterns
    author_pattern = re.compile(r'(?:by\s+|written\s+by\s+|author\s+)\s*([A-Z][a-zA-Z\s,]+)', re.IGNORECASE)
    year_pattern = re.compile(r'\b(\d{4})\b')
    title_pattern = re.compile(r'([A-Z][a-zA-Z\s]+(?:[a-zA-Z])+)')  # Refined title pattern
    
    # Attempt matching for author
    author_match = author_pattern.search(text)
    if author_match:
        bibliography_info['author'] = author_match.group(1).strip()
    
    # Attempt matching for year
    year_match = year_pattern.search(text)
    if year_match:
        bibliography_info['year'] = year_match.group(0).strip()
    
    # Attempt matching for title
    title_match = title_pattern.search(text)
    if title_match:
        bibliography_info['title'] = title_match.group(0).strip()

    # If no matches are found, attempt NLP processing for a better approach
    if not bibliography_info['author']:
        bibliography_info['author'] = extract_author_from_nlp(text)
    
    if not bibliography_info['year']:
        bibliography_info['year'] = extract_year_from_nlp(text)
    
    if not bibliography_info['title']:
        bibliography_info['title'] = extract_title_from_nlp(text)
    
    return bibliography_info

# Function to extract author using NLP
def extract_author_from_nlp(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return ''

# Function to extract year using NLP
def extract_year_from_nlp(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "DATE":
            return ent.text[-4:]  # Extract year
    return ''

# Function to extract title using NLP
def extract_title_from_nlp(text):
    doc = nlp(text)
    sentences = text.split('. ')
    # Heuristic: First sentence is often the title
    if len(sentences) > 0:
        return sentences[0]
    return ''

# Function to extract text from an image
def extract_text_from_image(image_bytes):
    try:
        image = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error extracting text from image: {e}")
        return ""

# Function to extract text from a PDF using PyMuPDF
def extract_text_from_pdf(pdf_bytes):
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

# Function to format bibliography information for download
def format_bibliography_info(bibliography_info):
    return f"Title: {bibliography_info['title']}\n\nAuthor: {bibliography_info['author']}\n\nYear: {bibliography_info['year']}"

# Function to download as Word
def download_word(content):
    doc = docx.Document()
    doc.add_paragraph(content)
    output_file = io.BytesIO()
    doc.save(output_file)
    output_file.seek(0)
    return output_file

# Function to download as PDF
def download_pdf(content):
    output_file = io.BytesIO()
    c = canvas.Canvas(output_file, pagesize=letter)
    text_object = c.beginText(40, 750)
    text_object.setFont("Helvetica", 12)
    lines = content.splitlines()
    for line in lines:
        text_object.textLine(line)
    c.drawText(text_object)
    c.showPage()
    c.save()
    output_file.seek(0)
    return output_file

# Function to download as Excel
def download_excel(content):
    df = pd.DataFrame({"Content": [content]})
    output_file = io.BytesIO()
    df.to_excel(output_file, index=False)
    output_file.seek(0)
    return output_file

# Initialize session state variables
if 'page' not in st.session_state:
    st.session_state['page'] = 'main'

if 'file_name' not in st.session_state:
    st.session_state['file_name'] = 'bibliography_content'

if 'theme' not in st.session_state:
    st.session_state['theme'] = 'light'

# Apply custom theme based on session state
theme_styles = {
    'light': """
    <style>
    .main-container {
        background-color: #f5f5f5;
        padding: 20px;
        border-radius: 10px;
    }
    .btn {
        background-color: #4CAF50;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        font-size: 16px;
    }
    .btn:hover {
        background-color: #45a049;
    }
    .header {
        font-size: 36px;
        font-weight: 600;
        color: #333;
    }
    .subheader {
        font-size: 24px;
        font-weight: 500;
        color: #555;
    }
    .text-box {
        border: 2px solid #ddd;
        padding: 10px;
        border-radius: 5px;
        background-color: #fff;
        color: black;
        max-height: 400px;
        overflow-y: auto;
        white-space: pre-wrap;
        font-family: monospace;
    }
    .file-name-container {
        display: flex;
        align-items: center;
        margin-bottom: 20px;
    }
    .file-name-container input {
        flex: 1;
        padding: 10px;
        border-radius: 5px;
        border: 1px solid #ccc;
    }
    .file-name-container button {
        margin-left: 10px;
    }
    </style>
    """,
    'dark': """
    <style>
    .main-container {
        background-color: #333;
        padding: 20px;
        border-radius: 10px;
        color: #f5f5f5;
    }
    .btn {
        background-color: #4CAF50;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        font-size: 16px;
    }
    .btn:hover {
        background-color: #45a049;
    }
    .header {
        font-size: 36px;
        font-weight: 600;
        color: #f5f5f5;
    }
    .subheader {
        font-size: 24px;
        font-weight: 500;
        color: #ddd;
    }
    .text-box {
        border: 2px solid #555;
        padding: 10px;
        border-radius: 5px;
        background-color: #444;
        color: #f5f5f5;
        max-height: 400px;
        overflow-y: auto;
        white-space: pre-wrap;
        font-family: monospace;
    }
    .file-name-container {
        display: flex;
        align-items: center;
        margin-bottom: 20px;
    }
    .file-name-container input {
        flex: 1;
        padding: 10px;
        border-radius: 5px;
        border: 1px solid #777;
    }
    .file-name-container button {
        margin-left: 10px;
    }
    </style>
    """
}

# Apply the selected theme style
st.markdown(theme_styles[st.session_state['theme']], unsafe_allow_html=True)

# Main and About page content
def main_page():
    st.title("Bibliography Extraction Tool")
    
    menu = ["Extract Bibliography", "Summarize Text"]
    choice = st.sidebar.selectbox("Choose an action", menu)
    
    if choice == "Extract Bibliography":
        st.subheader("Upload a file to extract bibliography information")

        uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx", "pptx", "jpg", "jpeg", "png"])
        if uploaded_file is not None:
            file_type = uploaded_file.type
            file_bytes = uploaded_file.read()

            if file_type == "image/jpeg" or file_type == "image/png" or file_type == "image/jpg":
                text = extract_text_from_image(file_bytes)
            elif file_type == "application/pdf":
                text = extract_text_from_pdf(file_bytes)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(io.BytesIO(file_bytes))
                text = "\n".join([para.text for para in doc.paragraphs])
            elif file_type == "application/vnd.ms-powerpoint":
                presentation = Presentation(io.BytesIO(file_bytes))
                text = "\n".join([slide.shapes.title.text if slide.shapes.title else "" for slide in presentation.slides])
            else:
                text = file_bytes.decode("utf-8")
            
            st.subheader("Extracted Text:")
            st.text_area("Text", value=text, height=300)

            bibliography_info = extract_bibliography_info(text)
            st.subheader("Extracted Bibliography Information:")
            st.write(f"Title: {bibliography_info['title']}")
            st.write(f"Author: {bibliography_info['author']}")
            st.write(f"Year: {bibliography_info['year']}")

            download_format = st.selectbox("Download Format", ["Word", "PDF", "PowerPoint", "Excel"])

            if download_format == "Word":
                word_file = download_word(format_bibliography_info(bibliography_info))
                st.download_button("Download Word File", word_file, file_name=f"{st.session_state['file_name']}.docx")
            elif download_format == "PDF":
                pdf_file = download_pdf(format_bibliography_info(bibliography_info))
                st.download_button("Download PDF File", pdf_file, file_name=f"{st.session_state['file_name']}.pdf")
            elif download_format == "Excel":
                excel_file = download_excel(format_bibliography_info(bibliography_info))
                st.download_button("Download Excel File", excel_file, file_name=f"{st.session_state['file_name']}.xlsx")
    
    elif choice == "Summarize Text":
        st.subheader("Enter text to summarize")
        text_input = st.text_area("Input Text", height=300)
        if st.button("Summarize"):
            summary = simple_summarize(text_input)
            st.subheader("Summary")
            st.write(summary)

# Display the main page
if st.session_state['page'] == 'main':
    main_page()
